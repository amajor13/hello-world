import os
import logging
from typing import List, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentLoader:
    """Handles loading and processing documents from various file formats."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document loader.
        
        Args:
            chunk_size: Size of text chunks for splitting
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def load_pdf(self, file_path: str) -> str:
        """Load text from a PDF file."""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            return ""
    
    def load_docx(self, file_path: str) -> str:
        """Load text from a DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            return ""
    
    def load_txt(self, file_path: str) -> str:
        """Load text from a TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            return ""
    
    def load_document(self, file_path: str) -> str:
        """Load a document based on its file extension."""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.load_pdf(file_path)
        elif file_extension == '.docx':
            return self.load_docx(file_path)
        elif file_extension in ['.txt', '.md']:
            return self.load_txt(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_extension}")
            return ""
    
    def load_documents_from_directory(self, directory_path: str) -> List[LangChainDocument]:
        """
        Load all supported documents from a directory and return chunked documents.
        
        Args:
            directory_path: Path to the directory containing documents
            
        Returns:
            List of LangChain Document objects with metadata
        """
        documents = []
        supported_extensions = ['.pdf', '.docx', '.txt', '.md']
        
        for file_path in Path(directory_path).rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                logger.info(f"Loading document: {file_path}")
                
                text = self.load_document(str(file_path))
                if text.strip():
                    # Split the text into chunks
                    chunks = self.text_splitter.split_text(text)
                    
                    # Create LangChain documents with metadata
                    for i, chunk in enumerate(chunks):
                        doc = LangChainDocument(
                            page_content=chunk,
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "chunk_index": i,
                                "total_chunks": len(chunks),
                                "file_type": file_path.suffix.lower()
                            }
                        )
                        documents.append(doc)
        
        logger.info(f"Loaded {len(documents)} document chunks from {directory_path}")
        return documents
    
    def load_single_document(self, file_path: str) -> List[LangChainDocument]:
        """
        Load a single document and return chunked documents.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of LangChain Document objects with metadata
        """
        documents = []
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File not found: {file_path}")
            return documents
        
        logger.info(f"Loading single document: {file_path}")
        text = self.load_document(file_path)
        
        if text.strip():
            chunks = self.text_splitter.split_text(text)
            
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "file_type": path.suffix.lower()
                    }
                )
                documents.append(doc)
        
        logger.info(f"Loaded {len(documents)} chunks from {file_path}")
        return documents